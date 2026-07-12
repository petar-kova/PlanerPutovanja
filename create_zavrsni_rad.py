#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Skript za kreiranje završnog rada – PlanerPutovanja
VSITE formatiranje: Times New Roman 12, Justify, 1.5 prored
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Margine ────────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.left_margin   = Cm(3)
section.bottom_margin = Cm(3)
section.right_margin  = Cm(2)
section.top_margin    = Cm(2)

# ── Pomoćne funkcije ────────────────────────────────────────────────────────────

def add_page_number(section):
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_para.clear()
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def set_body_para(para, size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=0):
    pf = para.paragraph_format
    pf.alignment = align
    pf.line_spacing = Pt(18)          # 1.5 × 12
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.first_line_indent = Cm(0)
    for run in para.runs:
        run.font.name  = 'Times New Roman'
        run.font.size  = Pt(size)
        run.font.bold  = bold
        run.font.italic = italic


def body(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         size=12, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = align
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.font.name   = 'Times New Roman'
    r.font.size   = Pt(size)
    r.font.bold   = bold
    r.font.italic = italic
    return p


def body_run(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_before=0, space_after=6):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = align
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.font.name   = 'Times New Roman'
        r.font.size   = Pt(size)
        r.font.bold   = bold
        r.font.italic = italic
    return p


def heading1(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment     = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing  = Pt(18)
    pf.space_before  = Pt(0)
    pf.space_after   = Pt(12)
    pf.page_break_before = True
    r = p.add_run(text.upper())
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.font.bold = True
    return p


def heading2(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment     = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing  = Pt(18)
    pf.space_before  = Pt(12)
    pf.space_after   = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.bold = True
    return p


def heading3(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment     = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing  = Pt(18)
    pf.space_before  = Pt(6)
    pf.space_after   = Pt(6)
    r = p.add_run(text)
    r.font.name  = 'Times New Roman'
    r.font.size  = Pt(12)
    r.font.bold  = True
    r.font.italic = True
    return p


def bullet(text, level=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.alignment     = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing  = Pt(18)
    p.paragraph_format.space_before  = Pt(0)
    p.paragraph_format.space_after   = Pt(3)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(10)
    return p


def caption_text(text, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = align
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.bold = False
    return p


def page_break():
    doc.add_page_break()


# ── Zaglavlje / Footer ──────────────────────────────────────────────────────────
add_page_number(section)

# ── SADRŽAJ (ručno) ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run('SADRŽAJ')
r.font.name = 'Times New Roman'
r.font.size = Pt(14)
r.font.bold = True
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(12)

toc_items = [
    ('1.', 'UVOD', '3'),
    ('2.', 'TEORIJSKE OSNOVE WEB APLIKACIJA I KORIŠTENIH TEHNOLOGIJA', '5'),
    ('2.1.', 'Web aplikacije i arhitektura klijent-poslužitelj', '5'),
    ('2.2.', 'ASP.NET Core MVC razvojni okvir', '6'),
    ('2.3.', 'Entity Framework Core i relacijske baze podataka', '8'),
    ('2.4.', 'Autentifikacija i autorizacija s ASP.NET Identity', '10'),
    ('2.5.', 'Vanjski API servisi', '11'),
    ('3.', 'DIZAJN I ARHITEKTURA APLIKACIJE ZA PLANIRANJE PUTOVANJA', '13'),
    ('3.1.', 'Zahtjevi i funkcionalnosti aplikacije', '13'),
    ('3.2.', 'Model podataka i shema baze podataka', '14'),
    ('3.3.', 'Arhitektura aplikacije', '17'),
    ('4.', 'PRAKTIČNI RAD – RAZVOJ WEB APLIKACIJE ZA PLANIRANJE PUTOVANJA', '19'),
    ('4.1.', 'Postavljanje projekta i konfiguracija okruženja', '19'),
    ('4.2.', 'Upravljanje putovanjima i destinacijama', '21'),
    ('4.3.', 'Praćenje aktivnosti i troškova', '24'),
    ('4.4.', 'Fotografski albumi', '26'),
    ('4.5.', 'Nadzorna ploča i statistike', '27'),
    ('4.6.', 'Generiranje PDF dokumenta', '29'),
    ('4.7.', 'Integracija vremenskih podataka i rute', '30'),
    ('4.8.', 'Kontakt forma', '31'),
    ('5.', 'ZAKLJUČAK', '33'),
    ('', 'LITERATURA', '35'),
    ('', 'SAŽETAK', '37'),
    ('', 'SUMMARY', '38'),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    tab_stops = p.paragraph_format.tab_stops
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    full = f'{num} {title}' if num else title
    r = p.add_run(f'{full}\t{page}')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    if not num or len(num) <= 2:
        r.font.bold = True

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# POPIS SLIKA
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
r = p.add_run('POPIS SLIKA')
r.font.name = 'Times New Roman'
r.font.size = Pt(14)
r.font.bold = True
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(12)

slike = [
    ('Slika 1.', 'Prikaz arhitekture MVC razvojnog obrasca', '7'),
    ('Slika 2.', 'Entity-Relationship dijagram baze podataka aplikacije PlanerPutovanja', '15'),
    ('Slika 3.', 'Arhitektura aplikacije – slojevi i ovisnosti', '18'),
    ('Slika 4.', 'Sučelje početne stranice aplikacije PlanerPutovanja', '20'),
    ('Slika 5.', 'Forma za kreiranje novog putovanja s višestrukim destinacijama', '22'),
    ('Slika 6.', 'Detalji putovanja – pregled aktivnosti, troškova i albuma', '25'),
    ('Slika 7.', 'Nadzorna ploča s grafikonima i statistikama putovanja', '28'),
    ('Slika 8.', 'Primjer generiranog PDF izvještaja putovanja', '30'),
]
for label, opis, str_br in slike:
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(3)
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(f'{label} {opis}\t{str_br}')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# POPIS TABLICA
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
r = p.add_run('POPIS TABLICA')
r.font.name = 'Times New Roman'
r.font.size = Pt(14)
r.font.bold = True
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(12)

tablice = [
    ('Tablica 1.', 'Pregled modela podataka aplikacije PlanerPutovanja', '14'),
    ('Tablica 2.', 'Pregled kontrolera i njihovih akcija', '17'),
    ('Tablica 3.', 'NuGet paketi korišteni u projektu', '19'),
    ('Tablica 4.', 'Načini prijevoza podržani u aplikaciji', '23'),
]
for label, opis, str_br in tablice:
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(3)
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(f'{label} {opis}\t{str_br}')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# POPIS KÔDOVA
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
r = p.add_run('POPIS KÔDOVA')
r.font.name = 'Times New Roman'
r.font.size = Pt(14)
r.font.bold = True
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(12)

kodovi = [
    ('Kôd 1.', 'Registracija servisa u Program.cs', '20'),
    ('Kôd 2.', 'Definicija modela Trip s anotacijama za validaciju', '21'),
    ('Kôd 3.', 'Metoda za filtriranje putovanja prema vremenskom statusu', '22'),
    ('Kôd 4.', 'Konfiguracija relacija u ApplicationDbContext', '16'),
    ('Kôd 5.', 'Servis za dohvat vremenskih podataka s OpenWeatherMap API-ja', '30'),
    ('Kôd 6.', 'Generiranje PDF dokumenta pomoću QuestPDF biblioteke', '29'),
]
for label, opis, str_br in kodovi:
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(3)
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(f'{label} {opis}\t{str_br}')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

# ══════════════════════════════════════════════════════════════════════════════
# 1. UVOD
# ══════════════════════════════════════════════════════════════════════════════
heading1('1. UVOD')

body(
    'Planiranje putovanja složen je proces koji uključuje odabir destinacija, organizaciju '
    'prijevoza, usklađivanje datuma, praćenje troškova i pohranjivanje uspomena. Sve te '
    'zadaće tradicionalno su se obavljale kombinacijom papirnatih bilježnica, proračunskih '
    'tablica i više nepovezanih internetskih usluga, što je korisnicima donosilo nepotrebnu '
    'složenost i gubitak vremena. Digitalizacija svakodnevnog života stvorila je potrebu za '
    'jedinstvenom, intuitivnom web aplikacijom koja na jednom mjestu objedinjuje sve korake '
    'organizacije putovanja.'
)

body(
    'Predmet ovog završnog rada je izrada web aplikacije '
    'PlanerPutovanja – sustava za planiranje i upravljanje putovanjima. Aplikacija '
    'korisnicima omogućuje kreiranje putovanja s više destinacija, planiranje aktivnosti, '
    'praćenje troškova u odnosu na zadani budžet, izradu fotografskih albuma i izvoz '
    'cjelovitog plana putovanja u PDF format. Cilj rada je prikazati kako se, primjenom '
    'modernih tehnologija i razvojnih obrazaca, može razviti funkcionalna i sigurna '
    'višekorisnička web aplikacija.'
)

body(
    'Motivacija za odabir ove teme leži u praktičnoj upotrebljivosti rješenja: aplikacija '
    'rješava svakodnevni problem koji pogađa velik broj korisnika, a pritom pruža '
    'zahtjevno tehničko okruženje za primjenu znanja stečenih tijekom studija. Kroz razvoj '
    'projekta primijenjeni su principi objektno orijentiranog programiranja (OOP), '
    'obrasca Model-Pogled-Kontroler (MVC), rada s relacijskim bazama podataka te '
    'integracije vanjskih programskih sučelja (API).'
)

body(
    'Teorijski dio rada, izložen u drugom poglavlju, opisuje ključne tehnologije i koncepte '
    'koji čine temelj aplikacije: arhitekturu web aplikacija, razvojni okvir ASP.NET Core, '
    'objektno-relacijsko preslikavanje putem Entity Framework Core te mehanizme '
    'autentifikacije i autorizacije. Treće poglavlje bavi se dizajnom i arhitekturom '
    'aplikacije – analizom zahtjeva, modelom podataka i organizacijom koda. Četvrto '
    'poglavlje prikazuje praktičnu realizaciju svake od funkcionalnosti sustava, uz '
    'odabrane isječke programskog koda i opis ključnih tehničkih odluka. Zaključak '
    'sažima postignute rezultate i navodi mogućnosti daljnjeg razvoja aplikacije.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. TEORIJSKE OSNOVE
# ══════════════════════════════════════════════════════════════════════════════
heading1('2. TEORIJSKE OSNOVE WEB APLIKACIJA I KORIŠTENIH TEHNOLOGIJA')

heading2('2.1. Web aplikacije i arhitektura klijent-poslužitelj')

body(
    'Web aplikacija je programska podrška kojoj se pristupa putem web preglednika i koja '
    'se izvršava na udaljenom poslužitelju, a rezultate prikazuje korisniku u obliku HTML '
    'stranica. Za razliku od klasičnih desktop aplikacija, web aplikacije ne zahtijevaju '
    'instalaciju i dostupne su s bilo kojeg uređaja koji ima pristup internetu.'
)

body(
    'Temeljna arhitektura web aplikacija zasniva se na modelu klijent-poslužitelj. Klijent '
    'je uobičajeno web preglednik koji šalje HTTP zahtjeve prema poslužitelju. Poslužitelj '
    'obrađuje zahtjev, izvodi poslovnu logiku, dohvaća podatke iz baze i vraća odgovor, '
    'najčešće u obliku HTML dokumenta, JSON-a ili drugog formata. Ovaj model dijeli '
    'odgovornosti između sučelja i logike obrade, čime se postiže skalabilnost i '
    'lakoća održavanja (Griffiths, Ian i sur., 2023).'
)

body(
    'Moderni web razvoj razlikuje dvije kategorije aplikacija: stranice renderiraneana '
    'poslužitelju (engl. Server-Side Rendering, SSR) i jednostranične aplikacije (engl. '
    'Single-Page Application, SPA). U SSR pristupu, koji je primijenjen u ovom radu, '
    'svaki korisnički zahtjev generira novi HTML odgovor na poslužitelju. Ovaj pristup '
    'jednostavniji je za implementaciju, prirodno podržava SEO optimizaciju i ne zahtijeva '
    'složene JavaScript okvire na strani klijenta.'
)

body(
    'Komunikacija između klijenta i poslužitelja odvija se putem protokola HTTP (engl. '
    'Hypertext Transfer Protocol). HTTP je protokol bez stanja (engl. stateless), što znači '
    'da svaki zahtjev nosi sve podatke potrebne za njegovu obradu. Upravljanje stanjem '
    'sesije postiže se mehanizmima poput kolačića (engl. cookies), JWT tokena ili '
    'serverskih sesija (Microsoft, 2024a).'
)

heading2('2.2. ASP.NET Core MVC razvojni okvir')

body(
    'ASP.NET Core je višeplatformski, modularni razvojni okvir tvrtke Microsoft namijenjen '
    'izgradnji modernih web aplikacija i servisa. Nasljeđuje prednosti prethodnih verzija '
    'ASP.NET platforme, ali je u potpunosti prerađen s naglaskom na performanse, '
    'prenosivost između operacijskih sustava i podršku za suvremene obrasce razvoja '
    '(Lock, Andrew, 2023).'
)

body(
    'Verzija .NET 8.0, korištena u ovom projektu, donosi značajna poboljšanja u '
    'brzini pokretanja aplikacije, smanjenju memorijskog otiska i optimizaciji '
    'kompajliranja. Razvojni okvir podržava ugrađenu injektciju ovisnosti (engl. '
    'Dependency Injection, DI) koja olakšava upravljanje životnim ciklusom servisa i '
    'smanjuje međuovisnost komponenti.'
)

body(
    'MVC (Model-View-Controller) arhitekturni je obrazac koji dijeli aplikaciju na tri '
    'sloja s jasno definiranim odgovornostima:'
)
bullet('Model (Model) – predstavlja podatke i poslovnu logiku aplikacije. Uključuje klase entiteta, validacijska pravila i repozitorij za pristup bazi podataka')
bullet('Pogled (View) – odgovoran je za prikaz podataka korisniku. U ASP.NET Core MVC-u pogledi su Razor datoteke (.cshtml) koje kombiniraju HTML markup s C# kodom')
bullet('Kontroler (Controller) – obrađuje korisničke zahtjeve, koordinira između modela i pogleda te vraća odgovarajući HTTP odgovor')

body(
    'Tok obrade zahtjeva u MVC obrascu prikazan je na Slici 1. Korisnikov zahtjev '
    'usmjerava se prema odgovarajućem kontroleru, kontroler dohvaća ili modificira '
    'podatke putem modela te prosljeđuje podatke pogledu koji generira HTML odgovor.'
)

caption_text('Slika 1. Prikaz arhitekture MVC razvojnog obrasca')

body(
    'ASP.NET Core MVC podržava atributnu validaciju modela, čime se pravila '
    'ispravnosti podataka definiraju izravno na svojstvima klase uz pomoć atributa '
    'poput [Required], [StringLength], [Range] i prilagođenih validatora. Okvir '
    'automatski validira model pri primitku zahtjeva i popunjava objekt ModelState '
    'eventulanim pogreškama, koje se mogu prikazati u pogledu uz pomoć Tag Helpera '
    '(Microsoft, 2024b).'
)

body(
    'Middleware pipeline mehanizam ASP.NET Core-a omogućuje modularno dodavanje '
    'funkcionalnosti u lanac obrade zahtjeva. Svaki middleware može obraditi zahtjev, '
    'proslijediti ga dalje ili prekinuti lanac. U aplikaciji PlanerPutovanja pipeline '
    'uključuje HTTPS preusmjeravanje, posluživanje statičkih datoteka, '
    'lokalizaciju, usmjeravanje, autentifikaciju i autorizaciju.'
)

heading2('2.3. Entity Framework Core i relacijske baze podataka')

body(
    'Entity Framework Core (EF Core) je objektno-relacijski mapper (ORM) koji '
    'programerima omogućuje rad s bazom podataka putem .NET objekata, bez pisanja SQL '
    'upita. Mapiranjem klasa entiteta na tablice baze podataka, EF Core prevodi LINQ '
    'upite u optimizirane SQL naredbe specifične za ciljni poslužitelj baze podataka '
    '(Microsoft, 2024c).'
)

body(
    'U arhitekturi aplikacije koristi se Code-First pristup: programer definira C# klase '
    'modela, a EF Core na temelju tih klasa generira shemu baze podataka putem migracija. '
    'Svaka promjena modela bilježi se kao zasebna migracija koja se može primijeniti ili '
    'poništiti, što omogućuje kontrolirano verzioniranje sheme baze.'
)

body(
    'Relacijska baza podataka Microsoft SQL Server korištena je kao backend za pohranu '
    'podataka. SQL Server pruža pouzdanost transakcijskog sustava, podršku za složene '
    'upite i dobru integraciju s Entity Framework Core-om. Veza s bazom podataka '
    'konfigurira se putem connection stringa u datoteci appsettings.json, što '
    'omogućuje jednostavnu promjenu baze bez modifikacije koda.'
)

body(
    'DbContext klasa središnji je element EF Core integracije. Klasa '
    'ApplicationDbContext naslijeđuje IdentityDbContext<User>, čime se integriraju '
    'ASP.NET Identity tablice, te definira DbSet kolekcije za sve entitete aplikacije. '
    'Relacije između entiteta konfiguriraju se u metodi OnModelCreating() korištenjem '
    'Fluent API-ja, što pruža veću kontrolu od atributnog pristupa:'
)

code_block(
'builder.Entity<Trip>()\n'
'    .HasMany(t => t.Activities)\n'
'    .WithOne(a => a.Trip)\n'
'    .HasForeignKey(a => a.TripId)\n'
'    .OnDelete(DeleteBehavior.Cascade);\n'
'\n'
'builder.Entity<Trip>()\n'
'    .HasMany(t => t.Albums)\n'
'    .WithOne(a => a.Trip)\n'
'    .HasForeignKey(a => a.TripId)\n'
'    .OnDelete(DeleteBehavior.Cascade);'
)
caption_text('Kôd 4. Konfiguracija relacija u ApplicationDbContext')

body(
    'Lazy loading, eager loading i explicit loading su tri strategije učitavanja '
    'povezanih entiteta u EF Core. U aplikaciji se pretežno koristi eager loading '
    'metodom Include(), kojom se uz osnovni entitet učitavaju i vezani entiteti u '
    'jednom upitu, smanjujući broj okruglih putovanja do baze (engl. N+1 problem).'
)

heading2('2.4. Autentifikacija i autorizacija s ASP.NET Identity')

body(
    'ASP.NET Core Identity je sustav za upravljanje korisničkim računima integriran '
    'u ASP.NET Core ekosustav. Pruža gotovu infrastrukturu za registraciju korisnika, '
    'prijavu, odjavljivanje, upravljanje lozinkama, uloge i tokene, čime programer '
    'izbjegava implementaciju osjetljivih sigurnosnih mehanizama od nule (Microsoft, 2024d).'
)

body(
    'Identity sustav pohranjuje korisnička pravila u SQL Server bazu putem EF Core-a. '
    'Klasa User, koja nasljeđuje IdentityUser, proširena je s kolekcijom putovanja '
    'kako bi se ostvarila veza između korisnika i njegovih podataka. Lozinke se pohranjuju '
    'kao bcrypt hash s nasumičnom soli, što ih štiti od napada riječnikom i brute-force '
    'tehnikama.'
)

body(
    'Autorizacija na razini kontrolera postiže se atributom [Authorize], koji '
    'zahtijeva da korisnik bude prijavljen za pristup zaštićenim resursima. Neautorizirani '
    'zahtjevi automatski se preusmjeravaju na stranicu za prijavu. Na razini podataka, '
    'primjenjuje se provjera vlasništva – svaki upit filtrira podatke prema '
    'CurrentUserId svojstvu izvučenom iz prijavljenog korisnika:'
)

code_block(
'private string CurrentUserId =>\n'
'    User.FindFirstValue(ClaimTypes.NameIdentifier) ?? string.Empty;\n'
'\n'
'var trips = _context.Trips\n'
'    .Where(t => t.UserId == CurrentUserId);'
)

body(
    'Ovaj pristup osigurava da korisnik može pristupati isključivo vlastitim podacima, '
    'što je temeljni sigurnosni zahtjev višekorisničke aplikacije.'
)

heading2('2.5. Vanjski API servisi')

body(
    'Moderna web aplikacija rijetko je potpuno samostalna – ona se integrira s vanjskim '
    'servisima koji pružaju specijalizirane funkcionalnosti. U aplikaciji PlanerPutovanja '
    'integrirani su dva vanjska API servisa: OpenWeatherMap za vremenske podatke i Google '
    'Maps Distance Matrix za izračun ruta.'
)

body(
    'API integracija realizirana je putem IHttpClientFactory mehanizma koji je dio '
    'ASP.NET Core dependency injection sustava. Ovaj pristup donosi nekoliko prednosti: '
    'upravljanje životnim ciklusom HttpClient instanci, konfiguracija specifičnih '
    'klijenata i primjena politika za ponavljanje zahtjeva (engl. retry policies). '
    'Svaki servis registriran je s vlastitim HTTP klijentom koji ima zadane parametre '
    'poput base adrese i timeouta.'
)

body(
    'OpenWeatherMap API pruža besplatni sloj (engl. free tier) koji dozvoljava do '
    '60 poziva po minuti. API vraća JSON odgovor s trenutnim vremenskim uvjetima za '
    'zadani grad, uključujući temperaturu, osjećaj temperature i opis stanja. Podaci '
    'se prikazuju na stranici detalja putovanja za svaku destinaciju.'
)

body(
    'Google Maps Distance Matrix API izračunava udaljenosti i trajanje vožnje između '
    'para lokacija. Primjenom memorijskog cachea (IMemoryCache) s vremenom isteka od '
    '10 minuta, broj API poziva svodi se na minimum, čime se kontroliraju troškovi '
    'korištenja naplatnog API-ja.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. DIZAJN I ARHITEKTURA
# ══════════════════════════════════════════════════════════════════════════════
heading1('3. DIZAJN I ARHITEKTURA APLIKACIJE ZA PLANIRANJE PUTOVANJA')

heading2('3.1. Zahtjevi i funkcionalnosti aplikacije')

body(
    'Analiza zahtjeva prvi je i ključni korak u razvoju programske podrške. '
    'Definiranjem funkcionalnih i nefunkcionalnih zahtjeva osigurava se da razvijeni '
    'sustav ispunjava potrebe korisnika i zadovoljava tehničke standarde kvalitete.'
)

body('Funkcionalni zahtjevi aplikacije PlanerPutovanja obuhvaćaju:')
bullet('registracija i prijava korisnika s provjerom identiteta')
bullet('kreiranje, uređivanje i brisanje putovanja s atributima naziva, destinacije, datuma, budžeta, valute i načina prijevoza')
bullet('dodavanje više destinacija u jedno putovanje uz definiranje redoslijeda obilaska i broja noćenja')
bullet('upravljanje aktivnostima vezanim za putovanje – unos naziva i bilješki')
bullet('evidencija troškova s usporedbom planiranog i stvarnog iznosa')
bullet('kreiranje fotografskih albuma s opisom, ocjenom i mogućnošću dodavanja više fotografija')
bullet('nadzorna ploča s ukupnim statistikama, grafom troškova i popisom nadolazećih putovanja')
bullet('izvoz plana putovanja u PDF format')
bullet('prikaz trenutnih vremenskih uvjeta za destinacije putovanja')
bullet('izračun udaljenosti i trajanja rute za putovanja automobilom')
bullet('filtriranje putovanja prema vremenskom statusu: sva, nadolazeća, prošla, aktualna')
bullet('kontakt forma za slanje poruka')

body('Nefunkcionalni zahtjevi uključuju:')
bullet('sigurnost: zaštita podataka korisnika, sprječavanje neovlaštenog pristupa tuđim podacima, zaštita od CSRF napada')
bullet('upotrebljivost: responzivno sučelje prilagođeno mobilnim uređajima, konzistentno vizualno oblikovanje')
bullet('performanse: primjena cachea za API pozive, učinkovito učitavanje relacija u EF Core-u')
bullet('održivost: modularna arhitektura zasnovana na MVC obrascu, jasna separacija odgovornosti')

heading2('3.2. Model podataka i shema baze podataka')

body(
    'Model podataka definira strukturu podataka kojima aplikacija upravlja. '
    'Prikazan je u Tablici 1 s pregledom svih entiteta, njihovih ključnih atributa '
    'i međusobnih veza.'
)

caption_text('Tablica 1. Pregled modela podataka aplikacije PlanerPutovanja')

from docx.shared import Inches
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Entitet'
hdr[1].text = 'Ključni atributi'
hdr[2].text = 'Veze'
hdr[3].text = 'Opis'
for cell in hdr:
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True

rows_data = [
    ('User', 'Id, Email, PasswordHash', '1:N s Trip', 'Korisnički račun (ASP.NET Identity)'),
    ('Trip', 'Id, Name, Destination, StartDate, EndDate, Budget, Currency, Transport', 'N:1 s User;\n1:N s Destinations, Activities, Expenses, Albums', 'Putovanje s osnovnim podacima'),
    ('TripDestination', 'Id, City, Nights, Order, TripId', 'N:1 s Trip', 'Pojedina destinacija unutar putovanja'),
    ('TripActivity', 'Id, Name, Notes, TripId', 'N:1 s Trip', 'Aktivnost planirana unutar putovanja'),
    ('Expense', 'Id, Name, Description, Amount, TripId', 'N:1 s Trip', 'Trošak evidentiran za putovanje'),
    ('TripAlbum', 'Id, TripId, Title, Review, Rating, CoverImagePath, CreatedAt', 'N:1 s Trip;\n1:N s TripPhoto', 'Fotografski album putovanja'),
    ('TripPhoto', 'Id, TripAlbumId, ImagePath, Caption, DisplayOrder, UploadedAt', 'N:1 s TripAlbum', 'Fotografija unutar albuma'),
    ('ContactMessage', 'Id, Name, Email, Subject, Message, SentAt, IsRead', '-', 'Poruka poslana putem kontakt forme'),
]

for row_data in rows_data:
    row = table.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        for para in row.cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

body('', space_before=6, space_after=6)

body(
    'Shema baze podataka prikazana je na Slici 2. Entiteti Trip, TripDestination, '
    'TripActivity, Expense, TripAlbum i TripPhoto međusobno su povezani stranim ključevima '
    's kaskadnim brisanjem, što znači da brisanjem putovanja automatski nestaju svi '
    'vezani zapisi.'
)

caption_text('Slika 2. Entity-Relationship dijagram baze podataka aplikacije PlanerPutovanja')

body(
    'Model Trip središnji je entitet sustava. Sadrži atribut Transport tipa enumeration '
    'koji definira šest načina prijevoza (automobil, avion, vlak, autobus, kruzer, ostalo), '
    'čime se podaci o načinu putovanja pohranjuju kao cjelobrojna vrijednost u bazi. '
    'Izvedeno svojstvo TotalNights automatski se izračunava iz razlike datuma završetka '
    'i početka putovanja bez pohrane u bazu (atribut [NotMapped]).'
)

body(
    'Posebna pozornost posvećena je validaciji podataka na razini modela. Svako svojstvo '
    'koje prima korisnički unos opremljeno je odgovarajućim anotacijama: [Required] '
    'osigurava nepraznost polja, [StringLength] ograničava duljinu teksta, a [Range] '
    'provjerava numeričke granice. Za složeniju validaciju ovisnu o drugom polju '
    '(datum završetka ne smije biti prije datuma početka) implementiran je prilagođeni '
    'validator CompareDatesAttribute koji nasljeđuje ValidationAttribute.'
)

heading2('3.3. Arhitektura aplikacije')

body(
    'Aplikacija PlanerPutovanja organizirana je prema MVC arhitekturnom obrascu, '
    'proširenom slojem servisa za enkapsulaciju složenije poslovne logike. Slika 3 '
    'prikazuje arhitekturu aplikacije s razvojem komponenti po slojevima.'
)

caption_text('Slika 3. Arhitektura aplikacije – slojevi i ovisnosti')

body(
    'Kontroleri su smješteni u mapu Controllers/ i odgovorni su za obradu HTTP zahtjeva. '
    'Svaki kontroler fokusiran je na jednu domenu: TripsController upravlja putovanjima, '
    'ExpensesController troškovima, TripActivitiesController aktivnostima, '
    'AlbumsController albumima i fotografijama, DashboardController nadzornom pločom, '
    'a HomeController statičkim stranicama i kontakt formom. Uz MVC kontrolere, '
    'RoutesApiController implementira JSON API endpoint za izračun ruta.'
)

caption_text('Tablica 2. Pregled kontrolera i njihovih akcija')

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Kontroler'
hdr2[1].text = 'Glavne akcije'
hdr2[2].text = 'Zahtijeva prijavu'
for cell in hdr2:
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True

rows2 = [
    ('HomeController', 'Index, Destinacije, Planer, Budzet, Galerija, Kontakt', 'Ne'),
    ('TripsController', 'Index, Details, Create, Edit, Delete, ExportPdf', 'Da'),
    ('ExpensesController', 'Create, Edit, Delete', 'Da'),
    ('TripActivitiesController', 'Create, Edit, Delete', 'Da'),
    ('AlbumsController', 'Index, Create, Details, Edit, Delete, AddPhotos, DeletePhoto', 'Da'),
    ('DashboardController', 'Index', 'Da'),
    ('RoutesApiController', 'Calculate (GET)', 'Ne'),
]
for rd in rows2:
    row = table2.add_row()
    for i, val in enumerate(rd):
        row.cells[i].text = val
        for para in row.cells[i].paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

body('', space_before=6, space_after=6)

body(
    'Servisni sloj smješten je u mapu Services/ i sadrži tri servisa: GoogleMapsService '
    'za izračun ruta, WeatherService za dohvat vremenskih podataka i TripPdfService '
    'za generiranje PDF dokumenta. Servisi su registrirani u DI spremniku '
    'kao scoped servisi i injektiraju se u kontrolere putem konstruktora, '
    'čime se ostvaruje labava sprega i testabilnost koda.'
)

body(
    'Pogledi su smješteni u mapu Views/ te organizirani u podmape prema kontrolerima. '
    'Zajednički elementi poput navigacijske trake, footera i skripti definirani su u '
    'zajedničkom layoutu Views/Shared/_Layout.cshtml koji se primjenjuje na sve '
    'stranice. Ovim pristupom eliminira se ponavljanje koda i osigurava konzistentnost '
    'vizualnog sučelja. Razor sintaksa omogućuje kombiniranje HTML-a s C# kodom '
    'unutar pogleda za uvjetno prikazivanje elemenata i petlje.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. PRAKTIČNI RAD
# ══════════════════════════════════════════════════════════════════════════════
heading1('4. PRAKTIČNI RAD – RAZVOJ WEB APLIKACIJE ZA PLANIRANJE PUTOVANJA')

heading2('4.1. Postavljanje projekta i konfiguracija okruženja')

body(
    'Projekt je razvijen u razvojnom okruženju Microsoft Visual Studio 2022 koristeći '
    '.NET 8.0 SDK. Kao predložak korišten je ASP.NET Core Web Application s MVC '
    'predloškom i opcijom za ASP.NET Core Identity koji generira početnu strukturu '
    'mapa i bazičnu konfiguraciju.'
)

caption_text('Tablica 3. NuGet paketi korišteni u projektu')

table3 = doc.add_table(rows=1, cols=3)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Paket'
hdr3[1].text = 'Verzija'
hdr3[2].text = 'Svrha'
for cell in hdr3:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True

rows3 = [
    ('Microsoft.EntityFrameworkCore.SqlServer', '8.x', 'EF Core provider za SQL Server'),
    ('Microsoft.EntityFrameworkCore.Tools', '8.x', 'Alati za migracije baze podataka'),
    ('Microsoft.AspNetCore.Identity.EntityFrameworkCore', '8.x', 'ASP.NET Identity integracija s EF Core'),
    ('QuestPDF', '2024.x', 'Generiranje PDF dokumenata'),
    ('Bootstrap', '5.3', 'CSS okvir za responzivno sučelje'),
]
for rd in rows3:
    row = table3.add_row()
    for i, val in enumerate(rd):
        row.cells[i].text = val
        for para in row.cells[i].paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

body('', space_before=6, space_after=6)

body(
    'Pokretanje aplikacije konfigurirano je u datoteci Program.cs koja redom registrira '
    'servise i gradi middleware pipeline. Datoteka appsettings.json pohranjuje '
    'konfiguraciju poput connection stringa za bazu podataka te API ključeva za '
    'OpenWeatherMap i Google Maps. Lokalizacija je postavljena na hr-HR kulturu '
    'kako bi se datumi i iznosi prikazivali prema hrvatskim standardima.'
)

code_block(
'builder.Services.AddDbContext<ApplicationDbContext>(options =>\n'
'    options.UseSqlServer(\n'
'        builder.Configuration.GetConnectionString("DefaultConnection")));\n'
'\n'
'builder.Services.AddDefaultIdentity<User>(options =>\n'
'{\n'
'    options.SignIn.RequireConfirmedAccount = false;\n'
'})\n'
'.AddEntityFrameworkStores<ApplicationDbContext>();\n'
'\n'
'builder.Services.AddScoped<GoogleMapsService>();\n'
'builder.Services.AddScoped<WeatherService>();\n'
'\n'
'QuestPDF.Settings.License = LicenseType.Community;'
)
caption_text('Kôd 1. Registracija servisa u Program.cs')

body(
    'Baza podataka inicijalizirana je putem EF Core migracija. Svaka migracija '
    'zabilježena je kao zasebna C# klasa u mapi Migrations/ i sadrži metode Up() za '
    'primjenu promjena i Down() za poništavanje. Ukupno je kreirano šest migracija '
    'koje prate evoluciju sheme od početne kreacije tablica do dodavanja albuma i '
    'kontakt poruka.'
)

body(
    'Početna stranica aplikacije (Slika 4) sadrži hero sekciju s kratkim opisom '
    'aplikacije, animiranom kartom koja vizualizira route s primjerom putovanja, '
    'kao i sekcije s pregledom mogućnosti i uputama za korištenje. Sučelje je '
    'implementirano uz Bootstrap 5 i prilagođeni CSS koji daje moderan izgled s '
    'gradijentima i animacijama.'
)

caption_text('Slika 4. Sučelje početne stranice aplikacije PlanerPutovanja')

heading2('4.2. Upravljanje putovanjima i destinacijama')

body(
    'Modul upravljanja putovanjima realiziran je unutar TripsController klase. '
    'Kontroler je zaštićen atributom [Authorize], što znači da sve akcije zahtijevaju '
    'prijavljenog korisnika. Metoda Index() dohvaća sva putovanja trenutnog '
    'korisnika uz eager loading aktivnosti, troškova i destinacija te podržava '
    'filtriranje prema vremenskom statusu putovanja:'
)

code_block(
'query = filter switch\n'
'{\n'
'    "upcoming" => query.Where(t => t.StartDate > today),\n'
'    "past"     => query.Where(t => t.EndDate < today),\n'
'    "current"  => query.Where(\n'
'        t => t.StartDate <= today && t.EndDate >= today),\n'
'    _          => query\n'
'};'
)
caption_text('Kôd 3. Metoda za filtriranje putovanja prema vremenskom statusu')

body(
    'Kreiranje putovanja odvija se kroz formu koja korisniku nudi unos naziva, '
    'destinacije, datuma polaska i povratka, budžeta, valute i načina prijevoza '
    '(Slika 5). Validacija se odvija u dva koraka: na strani klijenta uz pomoć '
    'Bootstrap validacijskih klasa i JavaScript provjere, te na strani poslužitelja '
    'provjером ModelState.IsValid nakon deserijalizacije forme.'
)

caption_text('Slika 5. Forma za kreiranje novog putovanja s višestrukim destinacijama')

body(
    'Posebnu pozornost zahtijevao je unos iznosa budžeta. Korisnici unose decimalne '
    'brojeve u različitim formatima ovisno o lokalnim konvencijama – s točkom ili '
    'zarezom kao separatorom decimala. Implementirana je metoda ParseBudget() koja '
    'normalizira unos u InvariantCulture format prije parsiranja, čime se '
    'izbjegavaju greške konverzije. Vrijednost se izvlači direktno iz Request.Form '
    'zbirke, zaobilazeći automatsko vezanje modela koje bi primijenilo kulturne '
    'postavke preglednika.'
)

body(
    'Višestruke destinacije implementirane su kao zasebni entitet TripDestination '
    'vezan uz putovanje. Svaka destinacija pohranjuje naziv grada, redni broj u itinerary '
    'redoslijedu i broj noćenja. Na strani sučelja, destinacije se dinamički dodaju i '
    'uklanjaju JavaScript kodom bez ponovnog učitavanja stranice, a šalju se '
    'poslužitelju kao niz indeksiranih polja forme.'
)

caption_text('Tablica 4. Načini prijevoza podržani u aplikaciji')

table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Table Grid'
hdr4 = table4.rows[0].cells
hdr4[0].text = 'Vrijednost enuma'
hdr4[1].text = 'Opis'
hdr4[2].text = 'Posebnost'
for cell in hdr4:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True

rows4 = [
    ('Car = 1', 'Automobil', 'Aktivira izračun rute i udaljenosti putem Google Maps API-ja'),
    ('Plane = 2', 'Avion', '-'),
    ('Train = 3', 'Vlak', '-'),
    ('Bus = 4', 'Autobus', '-'),
    ('CruiseShip = 5', 'Kruzer', 'Aktivira IsCruise oznaku za poseban prikaz'),
    ('Other = 6', 'Ostalo', '-'),
]
for rd in rows4:
    row = table4.add_row()
    for i, val in enumerate(rd):
        row.cells[i].text = val
        for para in row.cells[i].paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

body('', space_before=6, space_after=6)

body(
    'Model podataka Trip definiran je s odgovarajućim atributima koji osiguravaju '
    'ispravnost unesenih podataka. Prilagođeni validator CompareDatesAttribute '
    'uspoređuje datume i sprječava unos datuma završetka koji prethodi datumu početka:'
)

code_block(
'[Required(ErrorMessage = "Trip name is required.")]\n'
'[StringLength(100)]\n'
'public string Name { get; set; } = null!;\n'
'\n'
'[DataType(DataType.Date)]\n'
'[CompareDates("StartDate",\n'
'    ErrorMessage = "End date cannot be earlier than start date.")]\n'
'public DateTime EndDate { get; set; }\n'
'\n'
'[Column(TypeName = "decimal(18,2)")]\n'
'[Range(0, 1_000_000_000)]\n'
'public decimal? Budget { get; set; }'
)
caption_text('Kôd 2. Definicija modela Trip s anotacijama za validaciju')

heading2('4.3. Praćenje aktivnosti i troškova')

body(
    'Aktivnosti i troškovi vezani su uz konkretno putovanje i prikazuju se na stranici '
    'detalja putovanja (Slika 6). Stranica koristi eager loading s Include() i '
    'ThenInclude() metodama kako bi se u jednom upitu dohvatili putovanje, destinacije, '
    'aktivnosti, troškovi i albumi s fotografijama.'
)

caption_text('Slika 6. Detalji putovanja – pregled aktivnosti, troškova i albuma')

body(
    'Svaki trošak pohranjuje naziv, opis i iznos. Iznosi se zbrajaju i uspoređuju s '
    'planiranim budžetom putovanja, pri čemu se korisniku prikazuje ostatak budžeta. '
    'Ako je ukupna potrošnja veća od planiranog budžeta, sustav vizualno signalizira '
    'prekoračenje bojom i odgovarajućom porukom.'
)

body(
    'TripActivitiesController i ExpensesController koriste isti sigurnosni '
    'obrazac: svaka akcija koja dohvaća ili mijenja podatak prvo provjerava da '
    'li putovanje čijoj aktivnosti ili trošku se pristupa pripada trenutno '
    'prijavljenom korisniku. Ako provjera ne prođe, vraća se NotFound() odgovor '
    'koji korisniku ne otkriva postoji li traženi resurs ili mu samo nema pristup '
    '– ova tehnika sprječava enumeraciju tuđih podataka.'
)

heading2('4.4. Fotografski albumi')

body(
    'Modul fotografskih albuma implementiran je unutar AlbumsController klase. '
    'Album je vezan uz putovanje i može sadržavati neograničen broj fotografija '
    'uz naslov, recenziju putovanja i ocjenu od 1 do 5 zvjezdica.'
)

body(
    'Učitavanje fotografija odvija se kao multipart/form-data HTTP zahtjev. '
    'Na poslužitelju se svaka datoteka validira po tri kriterija: proširenje '
    'datoteke mora biti .jpg, .jpeg, .png ili .webp; veličina ne smije prelaziti '
    '5 MB; putanja za pohranu mora biti dostupna za pisanje. Prihvaćene fotografije '
    'pohranjuju se na datotečni sustav u mapu wwwroot/uploads/trips/{tripId}/ '
    's GUID generiranim imenom datoteke, čime se sprječava sukob naziva i '
    'onemogućuje predviđanje putanje od strane napadača.'
)

body(
    'Relativna putanja pohranjena u bazi podataka koristi se za generiranje URL-a '
    'kojim preglednik dohvaća sliku. Brisanje fotografije uklanja i fizičku '
    'datoteku s diska putem metode DeletePhysicalFile() koja robustno obrađuje '
    'IOException iznimku bez prekida toka izvršavanja.'
)

heading2('4.5. Nadzorna ploča i statistike')

body(
    'Nadzorna ploča (engl. dashboard) centralna je stranica koja prikazuje sažetak '
    'korisnikove aktivnosti (Slika 7). DashboardController agregira podatke iz '
    'baze i prosljeđuje ih u DashboardViewModel model pogleda.'
)

caption_text('Slika 7. Nadzorna ploča s grafikonima i statistikama putovanja')

body('Nadzorna ploča prikazuje sljedeće informacije:')
bullet('ukupan broj putovanja korisnika')
bullet('ukupna suma svih evidentiranih troškova')
bullet('ukupni planirani budžet i postotak iskorištenosti')
bullet('popis do pet nadolazećih putovanja sortiranih po datumu polaska')
bullet('grafikon troškova po putovanjima (do 12 putovanja)')
bullet('top 7 najposjećenijih destinacija s brojem posjeta')
bullet('popis putovanja kojima su troškovi dosegli 90 % ili više planiranog budžeta')

body(
    'Grafikoni su implementirani korištenjem Chart.js JavaScript biblioteke. '
    'Podaci se prenose iz kontrolera u pogled putem ViewBag svojstava te serijaliziraju '
    'u JSON format unutar Razor pogleda. Chart.js na temelju JSON podataka renderira '
    'interaktivne stupčaste i kružne grafikone u HTML Canvas elementima.'
)

body(
    'Posebno je naglašena logika identifikacije putovanja s prekoračenim budžetom. '
    'EF Core upit vrši se nad skupom svih putovanja korisnika koji imaju definiran '
    'budžet, a izračun postotka iskorištenosti provodi se u memoriji (LINQ to Objects) '
    'nakon što su podaci dohvaćeni iz baze. Putovanja čija iskorištenost prelazi '
    '90 % sortiraju se silazno prema postotku iskorištenosti.'
)

heading2('4.6. Generiranje PDF dokumenta')

body(
    'Funkcionalnost izvoza putovanja u PDF format realizirana je pomoću QuestPDF '
    'biblioteke otvorenog koda. QuestPDF koristi fluent API pattern koji opisuje '
    'izgled dokumenta kao hijerarhiju opisnih blokova: stranica, stupca, retka, '
    'tablice i teksta (QuestPDF, 2024).'
)

body(
    'Servis TripPdfService generira dokument koji sadrži naslovni blok s imenom '
    'putovanja i datumima, kartice s ključnim statistikama (trajanje, broj destinacija, '
    'broj aktivnosti, broj troškova), pregled budžeta s planiranim i ostvarenim iznosima, '
    'tablice destinacija, aktivnosti i troškova te sažetak s ostatkom budžeta (Slika 8). '
    'Generiranje PDF-a odvija se u memoriji i vraća se klijentu kao niz bajtova '
    'bez pohrane na disk, što smanjuje I/O opterećenje poslužitelja.'
)

caption_text('Slika 8. Primjer generiranog PDF izvještaja putovanja')

code_block(
'var document = Document.Create(container =>\n'
'{\n'
'    container.Page(page =>\n'
'    {\n'
'        page.Size(PageSizes.A4);\n'
'        page.Margin(24);\n'
'        page.Content().Column(column =>\n'
'        {\n'
'            column.Item().Element(c =>\n'
'                ComposeHero(c, tripName, itinerary,\n'
'                    description, statusText, statusColor));\n'
'            column.Item().Element(c =>\n'
'                ComposeDestinationsSection(c, destinations));\n'
'            column.Item().Element(c =>\n'
'                ComposeExpensesSection(c, expenses,\n'
'                    totalExpenses, remainingBudget));\n'
'        });\n'
'    });\n'
'});\n'
'return document.GeneratePdf();'
)
caption_text('Kôd 6. Generiranje PDF dokumenta pomoću QuestPDF biblioteke')

body(
    'Metoda za preuzimanje PDF-a (ExportPdf) dostupna je putem GET zahtjeva i '
    'vraća datoteku s Content-Disposition: attachment zaglavljem koje preglednikovim '
    'izaziva preuzimanje datoteke. Naziv datoteke generira se iz naziva putovanja, '
    'pri čemu se svi nevaljani znakovi za naziv datoteke zamjenjuju crticom.'
)

heading2('4.7. Integracija vremenskih podataka i rute')

body(
    'WeatherService dohvaća trenutne vremenske podatke za svaki grad destinacije '
    'putovanja. Servis koristi OpenWeatherMap Current Weather Data API koji prima '
    'naziv града i vraća JSON odgovor s temperaturom, osjećajem temperature i '
    'opisom stanja. Zahtjev se šalje s parametrom lang=hr što osigurava da su '
    'opisi vremenskog stanja na hrvatskom jeziku.'
)

code_block(
'var response = await client.GetAsync(\n'
'    $"weather?q={encodedCity}&units=metric&lang=hr&appid={apiKey}");\n'
'\n'
'if (!response.IsSuccessStatusCode) return null;\n'
'\n'
'var root = document.RootElement;\n'
'var temp = root.GetProperty("main")\n'
'               .GetProperty("temp").GetDecimal();\n'
'var desc = root.GetProperty("weather")[0]\n'
'               .GetProperty("description").GetString();'
)
caption_text('Kôd 5. Servis za dohvat vremenskih podataka s OpenWeatherMap API-ja')

body(
    'Na stranici detalja putovanja vremenski podaci prikazuju se za svaku destinaciju '
    'posebno. Kontroler iterira po listi gradova destinacija i za svaki grad '
    'paralelno (await unutar foreach petlje) dohvaća podatke i pohranjuje ih u '
    'Dictionary<string, WeatherInfo?> koji se prosljeđuje pogledu.'
)

body(
    'GoogleMapsService koristi Distance Matrix API za izračun ukupne kilometraže '
    'i trajanja rute putovanja automobilom. Servis prima listu lokacija (gradova '
    'destinacija) i sekvencijalno izračunava udaljenosti između susjednih lokacija '
    'koristeći origins i destinations parametre. Memorijski cache sprema rezultate '
    'upita na 10 minuta kako bi se smanjio broj naplatnih API poziva pri '
    'osvježavanju stranice.'
)

heading2('4.8. Kontakt forma')

body(
    'Kontakt forma implementirana je na javnoj stranici dostupnoj neregistriranim '
    'korisnicima. Model ContactMessage pohranjuje sve podatke poruke: ime, '
    'e-mail adresu pošiljatelja, temu i sadržaj poruke, zajedno s vremenskom oznakom '
    'slanja i oznakom o pročitanosti.'
)

body(
    'Obrada forme odvija se u HomeController u metodi Kontakt() s [HttpPost] atributom. '
    'Nakon validacije modela, poruka se pohranjuje u bazu podataka i korisnik '
    'se preusmjerava na istu stranicu. Poruka o uspješnom slanju prikazuje se '
    'putem TempData mehanizma koji prenosi podatke između preusmjeravanja, '
    'čime se implementira Post/Redirect/Get (PRG) obrazac koji sprječava '
    'dvostruko slanje forme osvježavanjem stranice.'
)

body(
    'Zaštita od lažnih zahtjeva (engl. Cross-Site Request Forgery, CSRF) '
    'implementirana je na svim POST formama putem [ValidateAntiForgeryToken] '
    'atributa i @Html.AntiForgeryToken() Razor helpera. ASP.NET Core automatski '
    'generira i validira CSRF token, čime se sprječava napad u kojemu zlonamjerna '
    'web stranica neovlašteno šalje zahtjeve u ime prijavljenog korisnika.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. ZAKLJUČAK
# ══════════════════════════════════════════════════════════════════════════════
heading1('5. ZAKLJUČAK')

body(
    'Izradom web aplikacije PlanerPutovanja ostvaren je cilj ovog završnog rada: '
    'razvijen je funkcionalan, siguran i upotrebljiv sustav za planiranje i upravljanje '
    'putovanjima koji objedinjuje sve korake organizacije putovanja na jednom mjestu. '
    'Aplikacija je realizirana korištenjem ASP.NET Core 8.0 MVC razvojnog okvira, '
    'Entity Framework Core ORM-a, SQL Server baze podataka i ASP.NET Identity '
    'sustava za autentifikaciju, čime su u praksi primijenjeni temeljni koncepti '
    'modernog web razvoja obrađeni u teorijskom dijelu rada.'
)

body(
    'Ostvarene su sve planirana funkcionalnosti: upravljanje putovanjima s podrškom '
    'za višestruke destinacije i filtriranje prema vremenskom statusu, evidencija '
    'aktivnosti i troškova s usporedbom planiranog budžeta, fotografski albumi '
    's pohranom i upravljanjem slikama, nadzorna ploča s interaktivnim grafikonima, '
    'izvoz plana putovanja u PDF dokument te integracija OpenWeatherMap i Google '
    'Maps API servisa za vremenske podatke i izračun rute. Posebna pažnja '
    'posvećena je sigurnosti – svaki pristup podacima filtriran je prema '
    'identitetu prijavljenog korisnika, a sve POST forme zaštićene su od CSRF napada.'
)

body(
    'Razvoj projekta potvrdio je praktičnu vrijednost MVC arhitekturnog obrasca: '
    'jasna podjela odgovornosti između modela, pogleda i kontrolera značajno je '
    'olakšala dodavanje novih funkcionalnosti bez narušavanja postojećeg koda. '
    'Servisni sloj dodatno je poboljšao modularnost, izdvajanjem složenije poslovne '
    'logike (generiranje PDF-a, API pozivi) iz kontrolera u zasebne klase.'
)

body(
    'Kao smjernice za daljnji razvoj aplikacije mogu se navesti: implementacija '
    'obavijesti e-poštom za nadolazeća putovanja, podrška za dijeljenje putovanja '
    's drugim korisnicima, mobilna aplikacija izgrađena na temelju REST API '
    'sučelja, sustav preporuka destinacija temeljen na korisnikovoj povijesti '
    'putovanja te napredna analitika troškova s kategorijama i vizualnim '
    'usporednim prikazima po putovanjima.'
)

# ══════════════════════════════════════════════════════════════════════════════
# LITERATURA
# ══════════════════════════════════════════════════════════════════════════════
heading1('LITERATURA')

lit_items = [
    'Griffiths, Ian, Lander, Matthew, (2023), Programming C# 12: Build Cloud, Web, and Desktop Applications, O\'Reilly Media',
    'Lock, Andrew, (2023), ASP.NET Core in Action, Third Edition, Manning Publications',
    'Microsoft, (2024a), Overview of ASP.NET Core, https://learn.microsoft.com/en-us/aspnet/core/introduction-to-aspnet-core (pristupljeno 1. 6. 2025.)',
    'Microsoft, (2024b), Model-View-Controller pattern in ASP.NET Core, https://learn.microsoft.com/en-us/aspnet/core/mvc/overview (pristupljeno 1. 6. 2025.)',
    'Microsoft, (2024c), Entity Framework Core documentation, https://learn.microsoft.com/en-us/ef/core/ (pristupljeno 3. 6. 2025.)',
    'Microsoft, (2024d), Introduction to Identity on ASP.NET Core, https://learn.microsoft.com/en-us/aspnet/core/security/authentication/identity (pristupljeno 3. 6. 2025.)',
    'OpenWeatherMap, Current weather data API, https://openweathermap.org/current (pristupljeno 5. 6. 2025.)',
    'QuestPDF, (2024), QuestPDF Documentation, https://www.questpdf.com/documentation/getting-started.html (pristupljeno 7. 6. 2025.)',
    'Google, Distance Matrix API documentation, https://developers.google.com/maps/documentation/distance-matrix/overview (pristupljeno 7. 6. 2025.)',
    'Troelsen, Andrew, Japikse, Philip, (2022), Pro C# 10 with .NET 6: Foundational Principles and Practices in Programming, Apress',
    'OWASP, OWASP Top Ten 2021, https://owasp.org/www-project-top-ten/ (pristupljeno 5. 6. 2025.)',
]

for item in lit_items:
    body(item, space_before=0, space_after=4)

# ══════════════════════════════════════════════════════════════════════════════
# SAŽETAK
# ══════════════════════════════════════════════════════════════════════════════
heading1('SAŽETAK')

body(
    'Ovaj završni rad opisuje razvoj web aplikacije PlanerPutovanja, višekorisničkog '
    'sustava za planiranje i upravljanje putovanjima realiziranog korištenjem '
    'tehnologije ASP.NET Core 8.0 MVC, Entity Framework Core i Microsoft SQL Server. '
    'Aplikacija korisnicima omogućuje kreiranje putovanja s višestrukim destinacijama, '
    'planiranje aktivnosti, praćenje troškova u odnosu na zadani budžet, '
    'izradu fotografskih albuma s pohranom slika, pregled statistika na nadzornoj '
    'ploči s interaktivnim grafikonima te izvoz plana putovanja u PDF dokument. '
    'Uz navedene funkcionalnosti, aplikacija integrira OpenWeatherMap API za prikaz '
    'trenutnih vremenskih uvjeta za destinacije putovanja te Google Maps Distance '
    'Matrix API za izračun udaljenosti i trajanja rute. Autentifikacija i autorizacija '
    'realizirane su uz pomoć ASP.NET Identity sustava koji osigurava sigurnu pohranu '
    'lozinki i izolaciju podataka između korisnika. U radu su detaljno opisani '
    'teorijski temelji korištenih tehnologija, dizajn modela podataka, arhitektura '
    'aplikacije te praktična implementacija svih funkcionalnih cjelina uz isječke '
    'programskog koda i opise tehničkih odluka.'
)

body(
    'Ključne riječi: web aplikacija, ASP.NET Core, MVC, Entity Framework Core, '
    'SQL Server, planiranje putovanja, REST API, PDF generiranje, autentifikacija'
)

# ══════════════════════════════════════════════════════════════════════════════
# SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
heading1('SUMMARY')

body(
    'This thesis describes the development of the PlanerPutovanja (Travel Planner) '
    'web application – a multi-user trip planning and management system built using '
    'ASP.NET Core 8.0 MVC, Entity Framework Core and Microsoft SQL Server. '
    'The application allows users to create trips with multiple destinations, plan '
    'activities, track expenses against a predefined budget, build photo albums with '
    'image storage, view statistics on a dashboard with interactive charts, and export '
    'the trip plan to a PDF document. Additionally, the application integrates the '
    'OpenWeatherMap API to display current weather conditions for trip destinations '
    'and the Google Maps Distance Matrix API to calculate driving distances and '
    'estimated travel times. Authentication and authorization are implemented using '
    'the ASP.NET Identity framework, which ensures secure password storage and '
    'data isolation between users. The thesis covers the theoretical foundations '
    'of the technologies used, the data model design, the application architecture '
    'and the practical implementation of all functional modules, supported by '
    'code snippets and descriptions of key technical decisions.'
)

body(
    'Keywords: web application, ASP.NET Core, MVC, Entity Framework Core, '
    'SQL Server, trip planning, REST API, PDF generation, authentication'
)

# ── Spremi dokument ─────────────────────────────────────────────────────────────
output_path = r'C:\Users\Student\Documents\Projekti\PlanerPutovanja\ZavrsniRad_PlanerPutovanja.docx'
doc.save(output_path)
print(f'Dokument uspješno kreiran: {output_path}')
